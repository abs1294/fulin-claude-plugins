# -*- coding: utf-8 -*-
"""
寓意科技報價單產生器。
用法：
    python build_quote.py <project.json> [-o 輸出檔.docx]

讀入一份專案資料 JSON（schema 見 SKILL.md），套用 assets/quote_template.docx
（已脫敏、保留寓意乙方資料與版式），產出格式一致的報價單 .docx。

設計重點：
- 明細列數「動態」：依 JSON 的 phases 自動增刪列（複製模板的階段列/項目列樣板）。
- 金額可「自動算」：每個項目給 md（人天），金額 = round(md * 含稅單價)；含稅單價 = unit_price*(1+tax)。
  也可在項目直接給 amount 覆寫。
- 總價自動加總；若 JSON 給了 total_incl_tax 則以它為準（並可校驗）。
"""
import sys, json, copy, argparse, os
import docx
from docx.oxml.ns import qn
from docx.table import _Row

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "..", "assets", "quote_template.docx")
FONT = "標楷體"   # 模板正文字型，附錄標題沿用以求一致

# 乙方固定資料（寓意科技）——已內建於模板右欄，此處僅供腳本備查，不需重填。

def set_cell(cell, lines):
    """整格替換文字，保留第一段/第一 run 的樣式；多行=多段。"""
    if isinstance(lines, str):
        lines = [lines]
    paras = cell.paragraphs
    rpr = paras[0].runs[0]._element.find(qn('w:rPr')) if paras[0].runs else None
    ppr = paras[0]._element.find(qn('w:pPr'))
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    for i, line in enumerate(lines):
        tp = p0 if i == 0 else cell.add_paragraph()
        if i > 0 and ppr is not None:
            tp._element.insert(0, copy.deepcopy(ppr))
        run = tp.add_run(str(line))
        if rpr is not None:
            run._element.insert(0, copy.deepcopy(rpr))

def money(n):
    return f"{int(round(n)):,}"

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("json_path")
    ap.add_argument("-o", "--out", default=None)
    args = ap.parse_args()

    with open(args.json_path, encoding="utf-8") as f:
        J = json.load(f)

    out = args.out or f"fable報價單_{J.get('client_short','客戶')}_{J.get('project_short','專案')}.docx"

    d = docx.Document(TEMPLATE)

    # ---------- 計價 ----------
    unit = J.get("unit_price", 10000)          # 未稅單價
    tax = J.get("tax_rate", 0.05)
    unit_incl = unit * (1 + tax)               # 含稅單價
    def item_amount(it):
        # amount 可為數字（含稅金額）或字串（如「額度內扣抵」，直接顯示、不計入加總）
        if "amount" in it:
            return it["amount"]
        if it.get("md") in (None, "", 0) and "md_text" in it:
            return ""                          # 無人天且僅文字→金額留空
        return round(it["md"] * unit_incl)

    # ---------- 表1 甲方 ----------
    c = J["client"]
    t1 = d.tables[0]
    left = [
        f"客戶名稱：{c.get('name','')}",
        f"客戶聯絡人：{c.get('contact','')}",
        f"Phone：{c.get('phone','')}",
        f"統一編號：{c.get('tax_id','')}",
        f"Email：{c.get('email','')}",
        f"地址：{c.get('address','')}",
    ]
    for i, row in enumerate(t1.rows):
        set_cell(row.cells[0], left[i])

    # 報價日期
    for p in d.paragraphs:
        if "報價日期" in p.text:
            rpr = p.runs[0]._element.find(qn('w:rPr')) if p.runs else None
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(f"報價日期：{J.get('quote_date','')}")
            if rpr is not None:
                run._element.insert(0, copy.deepcopy(rpr))
            break

    # ---------- 表2 概覽 + 明細 ----------
    t2 = d.tables[1]
    rows = t2.rows
    tbl = t2._tbl

    # 概覽
    for cell in rows[0].cells:
        set_cell(cell, J["project_name"])
    ov = J["overview"]
    set_cell(rows[1].cells[1], ov["manpower"])
    set_cell(rows[2].cells[1], ov["tech"])
    set_cell(rows[3].cells[1], ov["notes"])
    set_cell(rows[4].cells[1], ov["deliverables"])
    set_cell(rows[5].cells[1], ov.get("browsers", "支援 Google Chrome、Microsoft Edge、Safari、Firefox（建議使用最新版本）。"))
    set_cell(rows[6].cells[1], ov["schedule"])

    # 明細：模板列 9 = 表頭樣板、列 10 = 階段列樣板、列 11 = 項目列樣板
    header_tmpl = copy.deepcopy(rows[9]._tr)
    stage_tmpl = copy.deepcopy(rows[10]._tr)
    item_tmpl = copy.deepcopy(rows[11]._tr)
    tbl_tmpl = copy.deepcopy(tbl)   # 整表樣板（供附錄複製表格外框樣式）
    # 刪除模板明細列 10~29
    for idx in range(29, 9, -1):
        tbl.remove(rows[idx]._tr)

    total = 0
    anchor = t2.rows[9]._tr  # 「服務項目」表頭列
    for ph in J["phases"]:
        # 階段列
        tr = copy.deepcopy(stage_tmpl)
        anchor.addnext(tr); anchor = tr
        for cell in _Row(tr, t2).cells:
            set_cell(cell, ph["stage"])
        # 項目列
        for it in ph["items"]:
            amt = item_amount(it)
            if isinstance(amt, (int, float)) and amt != "":
                total += amt
                amt_disp = money(amt) if amt else ""
            else:
                amt_disp = str(amt)            # 字串金額（如「額度內扣抵」）直接顯示、不計入加總
            tr = copy.deepcopy(item_tmpl)
            anchor.addnext(tr); anchor = tr
            cells = _Row(tr, t2).cells
            set_cell(cells[0], it.get("no", ""))
            set_cell(cells[1], it["name"])
            set_cell(cells[2], it.get("weeks", ""))
            md_disp = f"{it['md']} 人天" if it.get("md") not in (None, "", 0) else it.get("md_text", "")
            set_cell(cells[3], md_disp)
            set_cell(cells[4], amt_disp)

    # 總計三列（模板尾端保留的 30/31/32，現為最後三列）
    total_incl = J.get("total_incl_tax", total)
    ar = t2.rows
    set_cell(ar[-3].cells[0], "總預算估計")
    for cell in ar[-3].cells[1:]:
        set_cell(cell, f"新台幣 {money(total_incl)} 元 (含稅)")
    set_cell(ar[-2].cells[0], "議價後折扣")
    disc = J.get("discount_text", "（依議定之計價模式與方案調整）")
    for cell in ar[-2].cells[1:]:
        set_cell(cell, disc)
    set_cell(ar[-1].cells[0], "報價有效期間")
    for cell in ar[-1].cells[1:]:
        set_cell(cell, f"至 {J.get('valid_until','')} 為止")

    # ---------- 表3 專案定義 ----------
    t3 = d.tables[2]
    set_cell(t3.rows[0].cells[1], J["project_definition"])

    # ---------- 表4 付款時程（可選覆寫）----------
    if "payment_schedule" in J:
        t4 = d.tables[3]
        set_cell(t4.rows[1].cells[1], J["payment_schedule"])

    # ---------- 附錄（選填）：另起新頁 + 標題 + 明細表 ----------
    if "appendix" in J:
        from docx.enum.text import WD_BREAK
        ap = J["appendix"]
        # 分頁符：附錄另起一頁
        pb = d.add_paragraph()
        pb.add_run().add_break(WD_BREAK.PAGE)
        # 標題段落——字型與正文一致（模板正文為標楷體）
        title_p = d.add_paragraph()
        run = title_p.add_run(ap.get("title", "附錄"))
        run.bold = True
        run.font.name = FONT
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            from docx.oxml import OxmlElement
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('ascii', 'eastAsia', 'hAnsi', 'cs'):
            rf.set(qn('w:'+a), FONT)
        # 建附錄表：複製整表樣板，清掉所有列，重建表頭+階段/項目列
        new_tbl = copy.deepcopy(tbl_tmpl)
        for tr in list(new_tbl.findall(qn('w:tr'))):
            new_tbl.remove(tr)
        # 表頭
        new_tbl.append(copy.deepcopy(header_tmpl))
        d.element.body.append(new_tbl)
        ap_tbl = d.tables[-1]
        for ph in ap.get("phases", []):
            tr = copy.deepcopy(stage_tmpl)
            new_tbl.append(tr)
            for cell in _Row(tr, ap_tbl).cells:
                set_cell(cell, ph["stage"])
            for it in ph.get("items", []):
                tr = copy.deepcopy(item_tmpl)
                new_tbl.append(tr)
                cells = _Row(tr, ap_tbl).cells
                set_cell(cells[0], it.get("no", ""))
                set_cell(cells[1], it["name"])
                set_cell(cells[2], it.get("weeks", ""))
                md_disp = f"{it['md']} 人天" if it.get("md") not in (None, "", 0) else it.get("md_text", "")
                set_cell(cells[3], md_disp)
                set_cell(cells[4], str(it.get("amount", "")))

    d.save(out)
    # 校驗
    warn = ""
    if "total_incl_tax" in J and abs(J["total_incl_tax"] - total) > 1:
        warn = f"  ⚠ 明細加總 {money(total)} 與指定總價 {money(J['total_incl_tax'])} 不一致，請確認"
    print(f"OK 產出：{out}")
    print(f"   明細加總（含稅）：{money(total)}{warn}")

if __name__ == "__main__":
    main()
